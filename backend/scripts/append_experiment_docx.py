"""Append JobGuard experiment results to an existing DOCX without deleting old text."""

from __future__ import annotations

import argparse
import json
import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


def shade(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_cell_width(cell, width: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    width_element = properties.find(qn("w:tcW"))
    if width_element is None:
        width_element = OxmlElement("w:tcW")
        properties.append(width_element)
    width_element.set(qn("w:w"), str(width))
    width_element.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    properties = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), "B8C4CE")
        borders.append(element)
    properties.append(borders)


def add_numbering(document: Document) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [int(value) for value in numbering.xpath("./w:abstractNum/@w:abstractNumId")]
    num_ids = [int(value) for value in numbering.xpath("./w:num/@w:numId")]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_format = OxmlElement("w:numFmt")
    num_format.set(qn("w:val"), "bullet")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•")
    paragraph_properties = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "420")
    tabs.append(tab)
    indentation = OxmlElement("w:ind")
    indentation.set(qn("w:left"), "420")
    indentation.set(qn("w:hanging"), "240")
    paragraph_properties.extend([tabs, indentation])
    level.extend([start, num_format, level_text, paragraph_properties])
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def set_run_font(run, name: str = "Microsoft YaHei", size: int = 10, bold: bool = False) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, size=14 if level == 1 else 12, bold=True)
    run.font.color.rgb = RGBColor(31, 78, 121)
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_with_next = True


def add_body(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run_font(first, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    paragraph.paragraph_format.space_after = Pt(5)


def add_bullets(doc: Document, items: list[str], num_id: int | None = None) -> None:
    """Render grouped short paragraphs when the source package has no numbering part."""
    for item in items:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(item)
        set_run_font(run)
        paragraph.paragraph_format.left_indent = Pt(12)
        paragraph.paragraph_format.space_after = Pt(3)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    usable = int(doc.sections[-1].page_width.twips - doc.sections[-1].left_margin.twips - doc.sections[-1].right_margin.twips)
    widths = widths or [usable // len(headers)] * len(headers)
    delta = usable - sum(widths)
    widths[-1] += delta
    table_width = table._tbl.tblPr.find(qn("w:tblW"))
    table_width.set(qn("w:w"), str(usable))
    table_width.set(qn("w:type"), "dxa")
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table._tbl.tblPr.append(layout)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    set_table_borders(table)
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    header_properties.append(repeat_header)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        set_run_font(run, size=9, bold=True)
        shade(cell, "D9EAF7")
        set_cell_width(cell, widths[index])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cells[index].text = ""
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(str(value))
            set_run_font(run, size=8)
            set_cell_width(cells[index], widths[index])
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("--matrix", type=Path, required=True)
    parser.add_argument("--backup", type=Path, required=True)
    args = parser.parse_args()

    if not args.backup.exists():
        shutil.copy2(args.input, args.backup)

    document = Document(args.input)
    original_paragraphs = [paragraph.text for paragraph in document.paragraphs]
    original_tables = [[cell.text for row in table.rows for cell in row.cells] for table in document.tables]
    matrix = json.loads(args.matrix.read_text(encoding="utf-8"))["rows"]
    bullet_num_id = None

    document.add_page_break()
    add_heading(document, "2026-08-25 JobGuard 实验追加记录", 1)
    add_body(document, "追加原则：前述 520 段原始终端记录完整保留；以下仅追加本轮已验证结果，不覆盖或删除旧实验。")

    add_heading(document, "1. 执行范围与环境", 2)
    add_bullets(document, [
        "模型：Qwen2.5-7B-Instruct，BF16。",
        "Serving：vLLM 0.7.3，OpenAI-compatible API。",
        "硬件：AutoDL RTX 4090 24GB；参数矩阵在同一实例、同一模型、同一并发负载下完成。",
        "代码同步：本地归档与服务器归档 SHA256 一致；实验产物已回传本地。",
    ], bullet_num_id)

    add_heading(document, "2. Agent Trace 与 Eval", 2)
    add_bullets(document, [
        "Trace 使用 request_id / trace_id / span_id 串联请求、Agent 节点与 LLM 调用；JSONL 追加写入。",
        "记录 provider、model、Prompt/Completion Token、延迟、异常、prompt_version、Cache 状态及脱敏 I/O 摘要。",
        "JobParser 真实链路 3 条请求：JSON 合法率、必填字段完整率、字段值准确率均为 100%，零错误。",
        "Agent 平均延迟 5437ms，P95 5860ms；Prompt Token 1602，Completion Token 719。",
    ], bullet_num_id)

    add_heading(document, "3. Agent 系统优化", 2)
    add_bullets(document, [
        "模型路由：短结构化任务优先本地 vLLM，复杂推理任务路由到推理模型，并记录路由原因与复杂度。",
        "Prompt 版本：岗位抽取 job_extract.v2、分类 job_classify.v1，随 Trace 写入。",
        "Cache：确定性请求使用 SHA256 精确匹配 TTL Cache，记录 hit / miss。",
        "并行：企业信息与口碑检索使用 asyncio.gather；批量岗位匹配限制为 4 路并发。",
        "Context：保留 System 消息与最近对话，在发送前按 16,000 字符预算压缩历史。",
        "测试：Agent 优化定向测试 10/10 通过；完整轻量测试 17/18，唯一失败是训练环境缺少 ChromaDB 导致导入失败。",
    ], bullet_num_id)

    add_heading(document, "4. vLLM 参数矩阵", 2)
    headers = ["配置", "显存MB", "RPS", "Tok/s", "TTFT P50", "TPOT P50", "成功率"]
    rows = []
    for row in matrix:
        rows.append([
            f'{row["max_model_len"]}/{row["gpu_memory_utilization"]:.2f}/' + ("on" if row["prefix_caching"] else "off"),
            str(row["gpu_memory_used_mb"]),
            f'{row["request_throughput_rps"]:.2f}', f'{row["output_throughput_tps"]:.2f}',
            f'{row["ttft_ms_p50"]:.2f}ms', f'{row["tpot_ms_p50"]:.2f}ms',
            f'{row["success_rate"] * 100:.0f}%',
        ])
    add_table(document, headers, rows, [1900, 1100, 850, 950, 1250, 1250, 850])
    add_body(document, "结论：短 Prompt 下各档差异较小。0.95 比 0.80 多占约 3.6GB 显存，但没有形成稳定吞吐优势；Prefix Cache 对本轮短且重复前缀的收益不稳定。建议默认 max_model_len=8192、gpu_memory_utilization=0.80，按业务是否存在长公共前缀决定是否开启 Prefix Cache。")

    add_heading(document, "5. 并发、Batch 与 KV Cache 既有有效结果", 2)
    add_table(document, ["并发", "RPS", "输出 Tok/s", "TTFT P50", "E2E P95"], [
        ["1", "0.75", "44.69", "47ms", "1.43s"],
        ["4", "2.89", "164.47", "47ms", "1.44s"],
        ["8", "5.55", "309.17", "58ms", "1.59s"],
        ["16", "9.53", "553.42", "65ms", "1.90s"],
    ])
    add_bullets(document, [
        "max_num_seqs=8 时，并发 16 吞吐降至 5.21 RPS，TTFT P50 升至约 1436ms，不建议设置过低。",
        "gpu_memory_utilization=0.70 / max_model_len=2048 时显存约 21.7GB；0.90 / 6000 时约 28.1GB（vGPU-32GB 实例既有结果）。",
    ], bullet_num_id)

    add_heading(document, "6. LoRA SFT", 2)
    add_table(document, ["LoRA r", "可训练参数", "Train Loss", "Val Loss", "字段准确率", "精确匹配率"], [
        ["4", "1,261,568", "0.2122", "0.1653", "83.85%", "0%"],
        ["8", "2,523,136", "0.1221", "0.0741", "86.15%", "0%"],
        ["16", "5,046,272", "0.0382", "0.0096", "98.46%", "80%"],
    ])
    add_body(document, "结论：选择 LoRA r=16 作为当前生产候选；相对 r=4/8 质量提升显著，额外显存成本很小。")

    add_heading(document, "7. DPO", 2)
    add_bullets(document, [
        "Preference 数据：80 Train / 20 Test；首轮过拟合结果保留作为失败对照。",
        "保守 DPO：学习率 1e-5、4 步；平均偏好 Margin 为 Base 0.389 → SFT 0.523 → DPO 0.598。",
        "DPO JSON 合法率 100%，字段准确率 93.85%，精确匹配率 50%。",
        "结论：DPO 提升偏好置信度，但任务质量低于 SFT r16，因此生产候选仍为 SFT r16。",
    ], bullet_num_id)

    add_heading(document, "8. 代码、数据与复现信息", 2)
    add_bullets(document, [
        "Agent Trace/Eval 与 vLLM 报告：backend/data/eval/。",
        "SFT/DPO 数据与报告：backend/data/post_training/。",
        "vLLM 矩阵：backend/data/eval/vllm_matrix_v2.json；逐配置日志位于 vllm_matrix_logs/。",
        "训练 Adapter 原路径：/root/autodl-tmp/jobguard_adapters/；本地保存 adapter_sha256.txt 校验清单。",
        "Commit 信息：当前工作区不含 .git 元数据，无法生成真实 commit；建议导入 Git 仓库后提交为 feat: add agent observability serving matrix and post-training eval。",
    ], bullet_num_id)

    document.save(args.input)
    verified = Document(args.input)
    assert [paragraph.text for paragraph in verified.paragraphs[: len(original_paragraphs)]] == original_paragraphs
    assert [[cell.text for row in table.rows for cell in row.cells] for table in verified.tables[: len(original_tables)]] == original_tables
    print(json.dumps({
        "original_paragraphs": len(original_paragraphs),
        "final_paragraphs": len(verified.paragraphs),
        "original_tables": len(original_tables),
        "final_tables": len(verified.tables),
        "prefix_preserved": True,
        "backup": str(args.backup),
    }, ensure_ascii=False))


if __name__ == "__main__":
    main()
