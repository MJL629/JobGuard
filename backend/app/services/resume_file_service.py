"""简历文件校验、文本提取与本地保存。"""

from __future__ import annotations

import hashlib
import io
import logging
import os
import re
import threading
import uuid
import zipfile
from dataclasses import dataclass
from pathlib import Path

from PIL import Image, UnidentifiedImageError

logger = logging.getLogger(__name__)

MAX_RESUME_BYTES = 10 * 1024 * 1024
MAX_PDF_PAGES = 20
MIN_EXTRACTED_CHARS = 20

_BACKEND_ROOT = Path(__file__).resolve().parents[2]
_UPLOAD_ROOT = _BACKEND_ROOT / "data" / "uploads"


class ResumeFileError(ValueError):
    """可安全返回给前端的简历文件错误。"""

    def __init__(self, message: str, *, code: str, status_code: int = 400):
        super().__init__(message)
        self.message = message
        self.code = code
        self.status_code = status_code


@dataclass
class ParsedResumeFile:
    original_name: str
    stored_name: str
    relative_path: str
    absolute_path: Path
    sha256: str
    media_type: str
    parser: str
    ocr_used: bool
    text: str
    content: bytes

    def public_metadata(self) -> dict:
        return {
            "original_name": self.original_name,
            "stored_name": self.stored_name,
            "sha256": self.sha256,
            "media_type": self.media_type,
            "parser": self.parser,
            "ocr_used": self.ocr_used,
            "extracted_chars": len(self.text),
        }


class ResumeFileService:
    """在不依赖扩展名的前提下解析常见简历文件。"""

    _ocr_engine = None
    _ocr_lock = threading.Lock()

    def prepare(self, user_id: int, filename: str | None, content: bytes) -> ParsedResumeFile:
        if not content:
            raise ResumeFileError("上传的文件为空，请重新选择简历文件", code="empty_file")
        if len(content) > MAX_RESUME_BYTES:
            raise ResumeFileError("简历文件不能超过 10MB", code="file_too_large", status_code=413)

        original_name = self._safe_original_name(filename)
        declared_suffix = Path(original_name).suffix.lower()
        detected_type, canonical_suffix = self._detect_type(content, declared_suffix)
        self._validate_suffix(declared_suffix, detected_type)

        try:
            text, parser, ocr_used = self._extract_text(content, detected_type)
        except ResumeFileError:
            raise
        except Exception as exc:
            logger.exception("简历文件文本提取失败，type=%s", detected_type)
            raise ResumeFileError(
                "文件可以读取，但文本提取失败。请尝试另存为 PDF 或 DOCX 后重新上传",
                code="text_extraction_failed",
                status_code=422,
            ) from exc

        text = self._normalize_text(text)
        if len(text) < MIN_EXTRACTED_CHARS:
            raise ResumeFileError(
                "没有识别到足够的简历文字。请确认图片清晰、PDF 未加密且文件内容完整",
                code="insufficient_text",
                status_code=422,
            )

        stored_name = f"{uuid.uuid4().hex}{canonical_suffix}"
        user_dir = (_UPLOAD_ROOT / str(user_id)).resolve()
        upload_root = _UPLOAD_ROOT.resolve()
        if upload_root not in user_dir.parents:
            raise ResumeFileError("无效的上传目录", code="invalid_storage_path", status_code=500)
        absolute_path = user_dir / stored_name
        relative_path = absolute_path.relative_to(_BACKEND_ROOT).as_posix()

        return ParsedResumeFile(
            original_name=original_name,
            stored_name=stored_name,
            relative_path=relative_path,
            absolute_path=absolute_path,
            sha256=hashlib.sha256(content).hexdigest(),
            media_type=detected_type,
            parser=parser,
            ocr_used=ocr_used,
            text=text,
            content=content,
        )

    @staticmethod
    def save(prepared: ParsedResumeFile) -> None:
        prepared.absolute_path.parent.mkdir(parents=True, exist_ok=True)
        temporary = prepared.absolute_path.with_suffix(prepared.absolute_path.suffix + ".part")
        try:
            with temporary.open("xb") as output:
                output.write(prepared.content)
                output.flush()
                os.fsync(output.fileno())
            temporary.replace(prepared.absolute_path)
        except Exception:
            if temporary.exists():
                temporary.unlink()
            raise

    @staticmethod
    def discard(prepared: ParsedResumeFile) -> None:
        """仅清理本次生成且仍位于上传目录内的文件。"""
        try:
            resolved = prepared.absolute_path.resolve()
            if _UPLOAD_ROOT.resolve() in resolved.parents and resolved.exists():
                resolved.unlink()
        except OSError:
            logger.warning("未能清理解析失败的上传文件：%s", prepared.stored_name)

    def _extract_text(self, content: bytes, detected_type: str) -> tuple[str, str, bool]:
        if detected_type == "application/pdf":
            return self._extract_pdf(content)
        if detected_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return self._extract_docx(content), "python-docx", False
        if detected_type.startswith("image/"):
            return self._ocr_image(content), "RapidOCR", True
        if detected_type == "text/plain":
            return self._decode_text(content), "text-decoder", False
        raise ResumeFileError("不支持该文件格式", code="unsupported_file_type")

    def _extract_pdf(self, content: bytes) -> tuple[str, str, bool]:
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(io.BytesIO(content), strict=False)
            if reader.is_encrypted:
                try:
                    if reader.decrypt("") == 0:
                        raise ResumeFileError("PDF 已加密，请先取消密码保护", code="encrypted_pdf")
                except ResumeFileError:
                    raise
                except Exception as exc:
                    raise ResumeFileError("PDF 已加密，请先取消密码保护", code="encrypted_pdf") from exc
            direct_text = "\n".join((page.extract_text() or "") for page in reader.pages)
        except ResumeFileError:
            raise
        except Exception as exc:
            raise ResumeFileError("PDF 文件已损坏或无法读取", code="invalid_pdf") from exc

        if len(self._normalize_text(direct_text)) >= MIN_EXTRACTED_CHARS:
            return direct_text, "PyPDF2", False

        try:
            import pypdfium2 as pdfium

            document = pdfium.PdfDocument(content)
            if len(document) > MAX_PDF_PAGES:
                raise ResumeFileError(
                    f"扫描版 PDF 最多支持 {MAX_PDF_PAGES} 页，请拆分后上传",
                    code="too_many_pdf_pages",
                )
            pages: list[str] = []
            for page_index in range(len(document)):
                page = document[page_index]
                image = page.render(scale=2.2).to_pil()
                buffer = io.BytesIO()
                image.convert("RGB").save(buffer, format="PNG")
                pages.append(self._ocr_image(buffer.getvalue()))
                page.close()
            document.close()
            return "\n".join(pages), "PyPDF2+PDFium+RapidOCR", True
        except ResumeFileError:
            raise
        except Exception as exc:
            raise ResumeFileError(
                "扫描版 PDF 的 OCR 识别失败，请上传更清晰的 PDF 或图片",
                code="pdf_ocr_failed",
                status_code=422,
            ) from exc

    @staticmethod
    def _extract_docx(content: bytes) -> str:
        try:
            from docx import Document

            document = Document(io.BytesIO(content))
            lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        lines.append("\t".join(cells))
            return "\n".join(lines)
        except Exception as exc:
            raise ResumeFileError("DOCX 文件已损坏或无法读取", code="invalid_docx") from exc

    @classmethod
    def _ocr_image(cls, content: bytes) -> str:
        try:
            with Image.open(io.BytesIO(content)) as image:
                image.verify()
        except (UnidentifiedImageError, OSError) as exc:
            raise ResumeFileError("图片文件已损坏或无法读取", code="invalid_image") from exc

        try:
            with cls._ocr_lock:
                if cls._ocr_engine is None:
                    from rapidocr import RapidOCR

                    cls._ocr_engine = RapidOCR()
                result = cls._ocr_engine(content)
            return "\n".join(result.txts or ())
        except Exception as exc:
            logger.exception("本地 OCR 执行失败")
            raise ResumeFileError(
                "图片 OCR 识别失败，请上传更清晰的图片或文字版 PDF",
                code="image_ocr_failed",
                status_code=422,
            ) from exc

    @staticmethod
    def _decode_text(content: bytes) -> str:
        if b"\x00" in content[:4096]:
            raise ResumeFileError("文件不是有效的文本简历", code="invalid_text_file")
        for encoding in ("utf-8-sig", "gb18030"):
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise ResumeFileError("文本编码无法识别，请保存为 UTF-8 后重新上传", code="invalid_text_encoding")

    @staticmethod
    def _detect_type(content: bytes, suffix: str) -> tuple[str, str]:
        if content.startswith(b"%PDF-"):
            return "application/pdf", ".pdf"
        if content.startswith(b"PK\x03\x04"):
            try:
                with zipfile.ZipFile(io.BytesIO(content)) as archive:
                    if "word/document.xml" in archive.namelist():
                        return (
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            ".docx",
                        )
            except zipfile.BadZipFile as exc:
                raise ResumeFileError("压缩格式文件已损坏", code="invalid_archive") from exc
            raise ResumeFileError("仅支持 DOCX，不支持普通 ZIP 文件", code="unsupported_archive")
        try:
            with Image.open(io.BytesIO(content)) as image:
                image_format = (image.format or "").upper()
                image.verify()
            image_types = {
                "PNG": ("image/png", ".png"),
                "JPEG": ("image/jpeg", ".jpg"),
                "WEBP": ("image/webp", ".webp"),
            }
            if image_format in image_types:
                return image_types[image_format]
        except (UnidentifiedImageError, OSError):
            pass
        if suffix in {".txt", ".md"}:
            return "text/plain", suffix
        raise ResumeFileError(
            "不支持该文件格式。请上传 PDF、DOCX、TXT、MD、PNG、JPG 或 WEBP",
            code="unsupported_file_type",
        )

    @staticmethod
    def _validate_suffix(suffix: str, detected_type: str) -> None:
        allowed = {
            "application/pdf": {".pdf"},
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": {".docx"},
            "image/png": {".png"},
            "image/jpeg": {".jpg", ".jpeg"},
            "image/webp": {".webp"},
            "text/plain": {".txt", ".md"},
        }
        if suffix not in allowed.get(detected_type, set()):
            raise ResumeFileError(
                "文件扩展名与实际内容不一致，请确认文件未被错误改名",
                code="file_type_mismatch",
            )

    @staticmethod
    def _safe_original_name(filename: str | None) -> str:
        name = Path((filename or "resume").replace("\\", "/")).name.strip()
        name = re.sub(r"[\x00-\x1f<>:\"/\\|?*]", "_", name)
        return name[:180] or "resume"

    @staticmethod
    def _normalize_text(text: str) -> str:
        text = text.replace("\x00", "").replace("\r\n", "\n").replace("\r", "\n")
        lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.split("\n")]
        return "\n".join(line for line in lines if line).strip()


resume_file_service = ResumeFileService()
