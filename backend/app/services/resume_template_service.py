"""Resume template catalog and deterministic DOCX rendering.

The four bundled templates are user-provided visual references.  Generated
documents use the same visual families without copying any sample identity or
resume facts from those files.
"""

from __future__ import annotations

import hashlib
import io
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT_DIR = Path(__file__).resolve().parents[2]
TEMPLATE_DIR = ROOT_DIR / "data" / "resume_templates"
OUTPUT_DIR = ROOT_DIR / "output" / "resumes"
USER_TEMPLATE_DIR = ROOT_DIR / "data" / "user_resume_templates"


@dataclass(frozen=True)
class TemplateSpec:
    id: str
    name: str
    description: str
    accent: str
    layout: str
    source_file: str
    editable: bool = True

    def as_dict(self) -> dict:
        return {
            "id": self.id,
            "name": self.name,
            "description": self.description,
            "accent": self.accent,
            "layout": self.layout,
            "source_file": self.source_file,
            "editable": self.editable,
        }


BUILTIN_TEMPLATES = (
    TemplateSpec(
        "template-01",
        "商务蓝分区",
        "蓝色标题带、信息密度高，适合通用职能和技术岗位。",
        "#4E7FA7",
        "blue_bands",
        "template-01.pdf",
    ),
    TemplateSpec(
        "template-02",
        "标准表格式",
        "信息表格清晰规整，适合偏正式、偏传统的投递场景。",
        "#111111",
        "formal_table",
        "template-02.docx",
    ),
    TemplateSpec(
        "template-03",
        "极简技术风",
        "黑白单栏、强调项目与技能，适合互联网和研发岗位。",
        "#111111",
        "minimal_lines",
        "template-03.docx",
    ),
    TemplateSpec(
        "template-04",
        "淡紫专业风",
        "淡紫分区、视觉层次较强，适合运营、商务和语言类岗位。",
        "#AAA7D5",
        "lavender_bands",
        "template-04.docx",
    ),
)


class ResumeTemplateService:
    def list_builtin(self) -> list[dict]:
        return [item.as_dict() for item in BUILTIN_TEMPLATES]

    def get_builtin(self, template_id: str | None) -> TemplateSpec | None:
        requested = template_id or "template-01"
        return next((item for item in BUILTIN_TEMPLATES if item.id == requested), None)

    def list_for_user(self, user_id: int) -> list[dict]:
        items = self.list_builtin()
        directory = USER_TEMPLATE_DIR / str(user_id)
        if directory.exists():
            for path in sorted(directory.glob("*.docx")):
                digest, _, original = path.stem.partition("__")
                items.append({
                    "id": f"custom:{digest}",
                    "name": original or "自定义模板",
                    "description": "用户上传的 DOCX 模板",
                    "accent": "#606266",
                    "layout": "custom_docx",
                    "source_file": path.name,
                    "editable": True,
                    "custom": True,
                })
        return items

    def save_custom(self, user_id: int, filename: str | None, content: bytes) -> dict:
        if not content or len(content) > 10 * 1024 * 1024:
            raise ValueError("模板文件为空或超过 10MB")
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as archive:
                if "word/document.xml" not in archive.namelist():
                    raise ValueError("仅支持有效的 DOCX 模板")
        except zipfile.BadZipFile as exc:
            raise ValueError("仅支持有效的 DOCX 模板") from exc
        safe_name = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff._-]+", "-", Path(filename or "自定义模板").stem)[:80]
        digest = hashlib.sha256(content).hexdigest()[:20]
        directory = USER_TEMPLATE_DIR / str(user_id)
        directory.mkdir(parents=True, exist_ok=True)
        path = directory / f"{digest}__{safe_name or '自定义模板'}.docx"
        if not path.exists():
            path.write_bytes(content)
        return {
            "id": f"custom:{digest}",
            "name": safe_name or "自定义模板",
            "description": "用户上传的 DOCX 模板",
            "accent": "#606266",
            "layout": "custom_docx",
            "source_file": path.name,
            "editable": True,
            "custom": True,
        }

    def get_custom_path(self, user_id: int, template_id: str) -> Path | None:
        if not template_id.startswith("custom:"):
            return None
        digest = template_id.split(":", 1)[1]
        if not re.fullmatch(r"[0-9a-f]{20}", digest):
            return None
        matches = list((USER_TEMPLATE_DIR / str(user_id)).glob(f"{digest}__*.docx"))
        return matches[0] if matches else None

    def render_docx(
        self,
        resume_id: int,
        markdown: str,
        template_id: str | None,
        user_id: int | None = None,
    ) -> str:
        spec = self.get_builtin(template_id)
        custom_path = self.get_custom_path(user_id, template_id or "") if user_id is not None else None
        if spec is None and custom_path is None:
            raise ValueError("未知的简历模板")
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        safe_template_id = (spec.id if spec else (template_id or "custom")).replace(":", "-")
        path = OUTPUT_DIR / f"resume_{resume_id}_{safe_template_id}.docx"
        document = Document(str(custom_path)) if custom_path else Document()
        if custom_path:
            body = document._element.body
            for child in list(body):
                if child.tag != qn("w:sectPr"):
                    body.remove(child)
        else:
            self._configure_document(document)
        blocks = list(self._parse_markdown(markdown))
        if spec and spec.layout == "formal_table":
            self._render_formal_table(document, blocks)
        else:
            render_spec = spec or TemplateSpec(
                template_id or "custom", "自定义模板", "", "#303133", "minimal_lines", custom_path.name
            )
            self._render_sections(document, blocks, render_spec)
        document.save(path)
        return str(path)

    @staticmethod
    def _configure_document(document: Document) -> None:
        section = document.sections[0]
        section.top_margin = Cm(1.35)
        section.bottom_margin = Cm(1.35)
        section.left_margin = Cm(1.55)
        section.right_margin = Cm(1.55)
        styles = document.styles
        styles["Normal"].font.name = "Microsoft YaHei"
        styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        styles["Normal"].font.size = Pt(10)

    def _render_sections(
        self, document: Document, blocks: list[tuple[str, str]], spec: TemplateSpec
    ) -> None:
        accent = self._hex_color(spec.accent)
        first_title = next((text for kind, text in blocks if kind == "h1"), "个人简历")
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT if spec.layout != "minimal_lines" else WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_after = Pt(10)
        run = title.add_run(first_title)
        run.bold = True
        run.font.size = Pt(28 if spec.layout != "minimal_lines" else 20)
        run.font.color.rgb = accent

        for kind, text in blocks:
            if kind == "h1":
                continue
            if kind == "h2":
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_before = Pt(8)
                paragraph.paragraph_format.space_after = Pt(5)
                run = paragraph.add_run(text)
                run.bold = True
                run.font.size = Pt(13)
                run.font.color.rgb = RGBColor(255, 255, 255) if spec.layout != "minimal_lines" else accent
                if spec.layout == "minimal_lines":
                    self._set_bottom_border(paragraph, accent)
                else:
                    self._shade_paragraph(paragraph, accent)
            elif kind == "h3":
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_before = Pt(4)
                paragraph.paragraph_format.space_after = Pt(2)
                run = paragraph.add_run(text)
                run.bold = True
                run.font.size = Pt(11)
            elif kind == "bullet":
                paragraph = document.add_paragraph(style="List Bullet")
                paragraph.paragraph_format.space_after = Pt(1)
                paragraph.add_run(text)
            else:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(3)
                paragraph.paragraph_format.line_spacing = 1.12
                paragraph.add_run(text)

    def _render_formal_table(
        self, document: Document, blocks: list[tuple[str, str]]
    ) -> None:
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(next((text for kind, text in blocks if kind == "h1"), "个人简历"))
        run.bold = True
        run.font.size = Pt(28)
        table = document.add_table(rows=0, cols=2)
        table.autofit = False
        table.columns[0].width = Cm(3.0)
        table.columns[1].width = Cm(14.5)
        current_section = "基本信息"
        pending: list[str] = []

        def flush() -> None:
            nonlocal pending
            if not pending:
                return
            row = table.add_row()
            row.cells[0].text = current_section
            row.cells[1].text = "\n".join(pending)
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                self._set_cell_border(cell)
            row.cells[0].paragraphs[0].runs[0].bold = True
            pending = []

        for kind, text in blocks:
            if kind == "h1":
                continue
            if kind == "h2":
                flush()
                current_section = text
            elif kind == "h3":
                pending.append(text)
            elif kind == "bullet":
                pending.append(f"• {text}")
            else:
                pending.append(text)
        flush()

    @staticmethod
    def _parse_markdown(markdown: str) -> Iterable[tuple[str, str]]:
        for raw_line in (markdown or "").splitlines():
            line = raw_line.strip()
            if not line:
                continue
            if line.startswith("### "):
                yield "h3", line[4:].strip()
            elif line.startswith("## "):
                yield "h2", line[3:].strip()
            elif line.startswith("# "):
                yield "h1", line[2:].strip()
            elif re.match(r"^[-*+]\s+", line):
                yield "bullet", re.sub(r"^[-*+]\s+", "", line)
            else:
                yield "text", re.sub(r"[*_`]", "", line)

    @staticmethod
    def _hex_color(value: str) -> RGBColor:
        value = value.lstrip("#")
        return RGBColor(*(int(value[i : i + 2], 16) for i in (0, 2, 4)))

    @staticmethod
    def _shade_paragraph(paragraph, color: RGBColor) -> None:
        p_pr = paragraph._p.get_or_add_pPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), str(color))
        p_pr.append(shading)

    @staticmethod
    def _set_bottom_border(paragraph, color: RGBColor) -> None:
        p_pr = paragraph._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "10")
        bottom.set(qn("w:color"), str(color))
        borders.append(bottom)
        p_pr.append(borders)

    @staticmethod
    def _set_cell_border(cell) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.first_child_found_in("w:tcBorders")
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            element = OxmlElement(f"w:{edge}")
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "6")
            element.set(qn("w:color"), "333333")
            borders.append(element)


resume_template_service = ResumeTemplateService()
