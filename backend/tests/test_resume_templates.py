from io import BytesIO
from pathlib import Path

from docx import Document

from app.services import resume_template_service as template_module
from app.services import resume_service as resume_module


SAMPLE_MARKDOWN = """# 张三

## 求职目标
后端开发工程师

## 专业技能
- Python
- FastAPI

## 项目经历
### 求职助手
- 负责接口开发与测试
"""


def test_all_builtin_templates_render_editable_docx(tmp_path, monkeypatch):
    monkeypatch.setattr(template_module, "OUTPUT_DIR", tmp_path)
    service = template_module.ResumeTemplateService()
    for index, item in enumerate(service.list_builtin(), 1):
        path = Path(service.render_docx(index, SAMPLE_MARKDOWN, item["id"], user_id=1))
        assert path.exists()
        document = Document(path)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        table_text = "\n".join(
            cell.text for table in document.tables for row in table.rows for cell in row.cells
        )
        assert "张三" in text + table_text
        assert "后端开发工程师" in text + table_text


def test_custom_docx_template_is_saved_listed_and_rendered(tmp_path, monkeypatch):
    custom_root = tmp_path / "custom"
    output_root = tmp_path / "output"
    monkeypatch.setattr(template_module, "USER_TEMPLATE_DIR", custom_root)
    monkeypatch.setattr(template_module, "OUTPUT_DIR", output_root)
    source = Document()
    source.add_paragraph("自定义模板占位内容")
    buffer = BytesIO()
    source.save(buffer)

    service = template_module.ResumeTemplateService()
    saved = service.save_custom(7, "我的模板.docx", buffer.getvalue())
    assert saved["id"].startswith("custom:")
    assert any(item["id"] == saved["id"] for item in service.list_for_user(7))

    rendered = Path(service.render_docx(9, SAMPLE_MARKDOWN, saved["id"], user_id=7))
    assert rendered.exists()
    result = Document(rendered)
    text = "\n".join(paragraph.text for paragraph in result.paragraphs)
    assert "张三" in text
    assert "自定义模板占位内容" not in text


def test_reportlab_fallback_exports_readable_chinese_pdf(tmp_path):
    pdf_path = tmp_path / "resume.pdf"

    exported = resume_module.ResumeService._export_pdf_via_reportlab(
        "# 张三\n\n## 项目经历\n### JobGuard\n- 使用 FastAPI 构建接口\n- 完成真实数据核验",
        str(pdf_path),
        "template-01",
    )

    assert exported is True
    assert pdf_path.read_bytes().startswith(b"%PDF-")
    assert pdf_path.stat().st_size > 1000
