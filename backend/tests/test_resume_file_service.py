import io

import pytest
from docx import Document
from PIL import Image

from app.services.resume_file_service import (
    MAX_RESUME_BYTES,
    ResumeFileError,
    ResumeFileService,
)


def test_utf8_text_resume_is_extracted():
    parsed = ResumeFileService().prepare(
        7,
        "我的简历.txt",
        "姓名：测试用户\n学校：测试大学\n专业：计算机科学\n技能：Python、MySQL".encode(),
    )
    assert "测试大学" in parsed.text
    assert parsed.parser == "text-decoder"
    assert parsed.ocr_used is False
    assert parsed.relative_path.startswith("data/uploads/7/")


def test_gb18030_text_resume_is_extracted():
    content = "姓名：测试用户\n五年工作经验\n毕业院校：测试大学\n专业：软件工程".encode("gb18030")
    parsed = ResumeFileService().prepare(7, "resume.txt", content)
    assert "软件工程" in parsed.text


def test_docx_paragraphs_and_tables_are_extracted():
    document = Document()
    document.add_paragraph("姓名：测试用户")
    document.add_paragraph("毕业院校：测试大学")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "专业"
    table.cell(0, 1).text = "计算机科学与技术"
    buffer = io.BytesIO()
    document.save(buffer)

    parsed = ResumeFileService().prepare(9, "resume.docx", buffer.getvalue())
    assert parsed.parser == "python-docx"
    assert "测试大学" in parsed.text
    assert "计算机科学与技术" in parsed.text


def test_image_uses_ocr_after_real_image_validation(monkeypatch):
    image = Image.new("RGB", (500, 160), "white")
    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    monkeypatch.setattr(
        ResumeFileService,
        "_ocr_image",
        classmethod(lambda cls, content: "姓名：测试用户\n学校：测试大学\n技能：Python 和 MySQL"),
    )

    parsed = ResumeFileService().prepare(10, "resume.png", buffer.getvalue())
    assert parsed.ocr_used is True
    assert parsed.parser == "RapidOCR"


def test_spoofed_extension_is_rejected():
    with pytest.raises(ResumeFileError) as error:
        ResumeFileService().prepare(1, "resume.pdf", "姓名：测试用户\n学校：测试大学".encode())
    assert error.value.code == "unsupported_file_type"


def test_actual_pdf_with_wrong_extension_is_rejected():
    with pytest.raises(ResumeFileError) as error:
        ResumeFileService().prepare(1, "resume.txt", b"%PDF-1.4\nnot a full pdf")
    assert error.value.code == "file_type_mismatch"


def test_too_large_file_is_rejected():
    with pytest.raises(ResumeFileError) as error:
        ResumeFileService().prepare(1, "resume.txt", b"a" * (MAX_RESUME_BYTES + 1))
    assert error.value.code == "file_too_large"


def test_invalid_image_is_rejected():
    with pytest.raises(ResumeFileError) as error:
        ResumeFileService().prepare(1, "resume.png", b"\x89PNG\r\n\x1a\ninvalid")
    assert error.value.code == "unsupported_file_type"
